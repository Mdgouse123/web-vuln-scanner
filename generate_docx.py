"""
generate_docx.py - Standalone script to convert a scan JSON report into a Word document.

Usage:
    python generate_docx.py reports/scan_example.com_20260709_103107.json
    python generate_docx.py reports/scan_example.com_20260709_103107.json --output my_report.docx
"""

import argparse
import json
import os
import sys
from datetime import datetime


def shade_cell(cell, hex_color="1e3a5f"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_heading(doc, text, level=1, r=30, g=58, b=138):
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(r, g, b)
    return p


def add_para(doc, text, bold=False, r=0, g=0, b=0, size=10):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(r, g, b)
    return p


def add_table(doc, headers, rows, col_widths=None):
    from docx.shared import Pt, RGBColor, Cm
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        shade_cell(cell, "1e3a5f")
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return t


SEV_COLORS = {
    "CRITICAL": (220, 38,  38),
    "HIGH":     (234, 88,  12),
    "MEDIUM":   (217, 119,  6),
    "LOW":      (8,  145, 178),
    "INFO":     (107, 114, 128),
}


def generate_docx(json_path, output_path=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Load JSON
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    url      = data.get("scan_target", "N/A")
    scan_time = data.get("scan_time", "N/A")
    summary  = data.get("summary", [])          # [{"severity": ..., "finding": ...}]
    details  = data.get("details", {})
    findings = [(s["severity"], s["finding"]) for s in summary]

    base     = details.get("base_info", {})
    ssl_info = details.get("ssl", {})
    hr       = details.get("headers", {})
    sf       = details.get("sensitive_files", {})
    sqli     = details.get("sqli", {})
    xss      = details.get("xss", {})
    redir    = details.get("open_redirect", {})

    # Output path
    if not output_path:
        base_name = os.path.splitext(os.path.basename(json_path))[0]
        output_path = os.path.join(os.path.dirname(json_path), base_name + ".docx")

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ------------------------------------------------------------------ title
    title = doc.add_heading("Web Vulnerability Scan Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 58, 95)

    meta = doc.add_paragraph(
        f"Target: {url}  |  Scan Time: {scan_time}  |  "
        f"Status: {base.get('status_code', 'N/A')}  |  "
        f"Server: {base.get('server', 'N/A')}"
    )
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph()

    # ------------------------------------------------------------------ 1. target info
    add_heading(doc, "1. Target Info")
    add_table(doc,
        ["Field", "Value"],
        [
            ["URL",           base.get("url", "N/A")],
            ["Final URL",     base.get("final_url", "N/A")],
            ["Status Code",   str(base.get("status_code", "N/A"))],
            ["Server",        base.get("server", "N/A")],
            ["Content-Type",  base.get("content_type", "N/A")],
            ["Response Time", f"{base.get('response_time_ms', 'N/A')} ms"],
        ],
        col_widths=[4, 12]
    )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 2. ssl
    add_heading(doc, "2. SSL / TLS Certificate")
    if ssl_info.get("error"):
        add_para(doc, f"[X] SSL Error: {ssl_info['error']}", bold=True, r=220, g=38, b=38)
    elif ssl_info.get("days_remaining") is None:
        add_para(doc, "Not applicable (HTTP).", r=100, g=100, b=100)
    else:
        days = ssl_info["days_remaining"]
        cr, cg, cb = (34, 197, 94) if days > 30 else ((250, 204, 21) if days > 7 else (220, 38, 38))
        add_para(doc, f"[OK] Certificate valid  |  Issuer: {ssl_info.get('ssl_issuer', 'N/A')}", bold=True, r=34, g=197, b=94)
        add_para(doc, f"Expires: {ssl_info.get('ssl_expiry', 'N/A')}  ({days} days remaining)", r=cr, g=cg, b=cb)
    doc.add_paragraph()

    # ------------------------------------------------------------------ 3. security headers
    add_heading(doc, "3. Security Headers")
    if "error" not in hr and hr:
        score     = hr.get("score", 0)
        max_score = hr.get("max_score", 1)
        pct       = int((score / max_score) * 100)
        cr, cg, cb = (34, 197, 94) if pct >= 70 else ((250, 204, 21) if pct >= 40 else (220, 38, 38))
        add_para(doc, f"Header Score: {score}/{max_score}  ({pct}%)", bold=True, r=cr, g=cg, b=cb)

        missing = hr.get("missing", [])
        if missing:
            add_para(doc, "Missing Headers:", bold=True, r=220, g=38, b=38)
            add_table(doc,
                ["Severity", "Header", "Recommendation"],
                [[h["severity"], h["header"], h["recommendation"]] for h in missing],
                col_widths=[2.5, 5, 9]
            )
            doc.add_paragraph()

        present = hr.get("present", [])
        if present:
            add_para(doc, "Present Headers:", bold=True, r=34, g=197, b=94)
            add_table(doc,
                ["Header", "Value"],
                [[h["header"], h["value"][:80]] for h in present],
                col_widths=[5, 11.5]
            )
            doc.add_paragraph()

        disc = hr.get("discouraged", [])
        if disc:
            add_para(doc, "Information Disclosure Headers:", bold=True, r=250, g=204, b=21)
            add_table(doc,
                ["Header", "Value", "Recommendation"],
                [[h["header"], h["value"], h["recommendation"]] for h in disc],
                col_widths=[3, 4, 9.5]
            )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 4. sensitive files
    add_heading(doc, "4. Sensitive Files & Directories")
    found_files = sf.get("found", [])
    cr, cg, cb  = (220, 38, 38) if found_files else (34, 197, 94)
    add_para(doc,
        f"Probed {sf.get('checked', 0)} paths - found {len(found_files)} exposed.",
        bold=True, r=cr, g=cg, b=cb
    )
    if found_files:
        add_table(doc,
            ["Severity", "Path", "Status", "Description"],
            [[f["severity"], f["path"], str(f["status_code"]), f["description"][:55]] for f in found_files],
            col_widths=[2.2, 4, 1.8, 8.5]
        )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 5. sqli
    add_heading(doc, "5. SQL Injection Detection")
    if not sqli.get("has_params"):
        add_para(doc, "No query parameters in URL. Skipping.", r=100, g=100, b=100)
    elif not sqli.get("vulnerable_params"):
        add_para(doc, "[OK] No SQL injection errors detected.", bold=True, r=34, g=197, b=94)
    else:
        add_table(doc,
            ["Parameter", "Payload", "Matched Signature"],
            [[v["param"], v["payload"], v["matched_signature"]] for v in sqli["vulnerable_params"]],
            col_widths=[3, 6, 7.5]
        )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 6. xss
    add_heading(doc, "6. XSS Detection")
    if not xss.get("has_params"):
        add_para(doc, "No query parameters in URL. Skipping.", r=100, g=100, b=100)
    elif not xss.get("vulnerable_params"):
        add_para(doc, "[OK] No reflected XSS detected.", bold=True, r=34, g=197, b=94)
    else:
        add_table(doc,
            ["Parameter", "Payload"],
            [[v["param"], v["payload"]] for v in xss["vulnerable_params"]],
            col_widths=[4, 12.5]
        )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 7. open redirect
    add_heading(doc, "7. Open Redirect Detection")
    if not redir.get("vulnerable_params"):
        add_para(doc, "[OK] No open redirect vulnerabilities detected.", bold=True, r=34, g=197, b=94)
    else:
        add_table(doc,
            ["Parameter", "Payload"],
            [[v["param"], v["payload"]] for v in redir["vulnerable_params"]],
            col_widths=[4, 12.5]
        )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 8. summary
    add_heading(doc, "8. Scan Summary")
    if not findings:
        add_para(doc, "[OK] No significant vulnerabilities detected!", bold=True, r=34, g=197, b=94)
    else:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        for i, h in enumerate(["Severity", "Finding"]):
            cell = t.rows[0].cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            shade_cell(cell, "1e3a5f")
        for sev, desc in findings:
            row = t.add_row()
            sc = row.cells[0]
            sc.text = sev
            rr, rg, rb = SEV_COLORS.get(sev, (0, 0, 0))
            run = sc.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(rr, rg, rb)
            dc = row.cells[1]
            dc.text = desc
            dc.paragraphs[0].runs[0].font.size = Pt(9)
        for row in t.rows:
            row.cells[0].width = Cm(3)
            row.cells[1].width = Cm(13.5)
        doc.add_paragraph()
        add_para(doc, f"Total findings: {len(findings)}", bold=True, r=220, g=38, b=38, size=11)

    # ------------------------------------------------------------------ footer
    doc.add_paragraph()
    doc.add_paragraph("-" * 80)
    footer = doc.add_paragraph(
        "Generated by WebVulnScanner v1.0 - For authorized security testing only."
    )
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Convert a scan JSON report into a Word (.docx) document.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python generate_docx.py reports/scan_example.com_20260709_103107.json
  python generate_docx.py reports/scan_example.com_20260709_103107.json --output report.docx
        """
    )
    parser.add_argument("json_file", help="Path to the scan JSON report file")
    parser.add_argument("--output", default=None, help="Output .docx file path (optional)")
    args = parser.parse_args()

    if not os.path.exists(args.json_file):
        print(f"[ERROR] File not found: {args.json_file}")
        sys.exit(1)

    try:
        from docx import Document
    except ImportError:
        print("[ERROR] python-docx not installed. Run: pip install python-docx")
        sys.exit(1)

    print(f"[*] Reading: {args.json_file}")
    out = generate_docx(args.json_file, args.output)
    print(f"[OK] Word document saved: {out}")


if __name__ == "__main__":
    main()
