"""
report_docx.py - Export scan results as a Word (.docx) document.
"""

import os
from datetime import datetime
from urllib.parse import urlparse


def export_docx(url, all_results, findings, normalize_url_fn):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    os.makedirs("reports", exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    host = urlparse(normalize_url_fn(url)).hostname or "unknown"
    path = f"reports/scan_{host}_{ts}.docx"

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ------------------------------------------------------------------ helpers

    def add_heading(text, level=1, r=30, g=58, b=138):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(r, g, b)
        return p

    def add_para(text, bold=False, r=0, g=0, b=0, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(r, g, b)
        return p

    def shade_cell(cell, hex_color="1e3a5f"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    def add_table(headers, rows, col_widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        hrow = t.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            shade_cell(cell, "1e3a5f")
        for row_data in rows:
            row = t.add_row()
            for i, val in enumerate(row_data):
                cell = row.cells[i]
                cell.text = str(val)
                cell.paragraphs[0].runs[0].font.size = Pt(9)
        if col_widths:
            for row in t.rows:
                for i, w in enumerate(col_widths):
                    if i < len(row.cells):
                        row.cells[i].width = Cm(w)
        return t

    SEV_COLORS = {
        "CRITICAL": (220, 38,  38),
        "HIGH":     (234, 88,  12),
        "MEDIUM":   (217, 119,  6),
        "LOW":      (8,  145, 178),
        "INFO":     (107, 114, 128),
    }

    base     = all_results.get("base_info", {})
    ssl_info = all_results.get("ssl", {})

    # ------------------------------------------------------------------ title
    title = doc.add_heading("Web Vulnerability Scan Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 58, 95)

    meta = doc.add_paragraph(
        f"Target: {url}  |  "
        f"Scan Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  "
        f"Status: {base.get('status_code', 'N/A')}  |  "
        f"Server: {base.get('server', 'N/A')}"
    )
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph()

    # ------------------------------------------------------------------ 1. target info
    add_heading("1. Target Info")
    add_table(
        ["Field", "Value"],
        [
            ["URL",           base.get("url", "N/A")],
            ["Final URL",     base.get("final_url", "N/A")],
            ["Status Code",   str(base.get("status_code", "N/A"))],
            ["Server",        base.get("server", "N/A")],
            ["Content-Type",  base.get("content_type", "N/A")],
            ["Response Time", f"{base.get('response_time_ms', 'N/A')} ms"],
        ],
        col_widths=[4, 12]
    )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 2. ssl
    add_heading("2. SSL / TLS Certificate")
    if ssl_info.get("error"):
        add_para(f"[X] SSL Error: {ssl_info['error']}", bold=True, r=220, g=38, b=38)
    elif ssl_info.get("days_remaining") is None:
        add_para("Not applicable (HTTP).", r=100, g=100, b=100)
    else:
        days = ssl_info["days_remaining"]
        cr, cg, cb = (34, 197, 94) if days > 30 else ((250, 204, 21) if days > 7 else (220, 38, 38))
        add_para(f"[OK] Certificate valid  |  Issuer: {ssl_info.get('ssl_issuer', 'N/A')}", bold=True, r=34, g=197, b=94)
        add_para(f"Expires: {ssl_info.get('ssl_expiry', 'N/A')}  ({days} days remaining)", r=cr, g=cg, b=cb)
    doc.add_paragraph()

    # ------------------------------------------------------------------ 3. headers
    add_heading("3. Security Headers")
    hr = all_results.get("headers", {})
    if "error" not in hr:
        score, max_score = hr.get("score", 0), hr.get("max_score", 1)
        pct = int((score / max_score) * 100)
        cr, cg, cb = (34, 197, 94) if pct >= 70 else ((250, 204, 21) if pct >= 40 else (220, 38, 38))
        add_para(f"Header Score: {score}/{max_score}  ({pct}%)", bold=True, r=cr, g=cg, b=cb)

        missing = hr.get("missing", [])
        if missing:
            add_para("Missing Headers:", bold=True, r=220, g=38, b=38)
            add_table(
                ["Severity", "Header", "Recommendation"],
                [[h["severity"], h["header"], h["recommendation"]] for h in missing],
                col_widths=[2.5, 5, 9]
            )
            doc.add_paragraph()

        present = hr.get("present", [])
        if present:
            add_para("Present Headers:", bold=True, r=34, g=197, b=94)
            add_table(
                ["Header", "Value"],
                [[h["header"], h["value"][:80]] for h in present],
                col_widths=[5, 11.5]
            )
            doc.add_paragraph()

        disc = hr.get("discouraged", [])
        if disc:
            add_para("Information Disclosure Headers:", bold=True, r=250, g=204, b=21)
            add_table(
                ["Header", "Value", "Recommendation"],
                [[h["header"], h["value"], h["recommendation"]] for h in disc],
                col_widths=[3, 4, 9.5]
            )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 4. sensitive files
    add_heading("4. Sensitive Files & Directories")
    sf = all_results.get("sensitive_files", {})
    found_files = sf.get("found", [])
    cr, cg, cb = (220, 38, 38) if found_files else (34, 197, 94)
    add_para(
        f"Probed {sf.get('checked', 0)} paths - found {len(found_files)} exposed.",
        bold=True, r=cr, g=cg, b=cb
    )
    if found_files:
        add_table(
            ["Severity", "Path", "Status", "Description"],
            [[f["severity"], f["path"], str(f["status_code"]), f["description"][:55]] for f in found_files],
            col_widths=[2.2, 4, 1.8, 8.5]
        )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 5. sqli
    add_heading("5. SQL Injection Detection")
    sqli = all_results.get("sqli", {})
    if not sqli.get("has_params"):
        add_para("No query parameters in URL. Skipping.", r=100, g=100, b=100)
    elif not sqli.get("vulnerable_params"):
        add_para("[OK] No SQL injection errors detected.", bold=True, r=34, g=197, b=94)
    else:
        add_table(
            ["Parameter", "Payload", "Matched Signature"],
            [[v["param"], v["payload"], v["matched_signature"]] for v in sqli["vulnerable_params"]],
            col_widths=[3, 6, 7.5]
        )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 6. xss
    add_heading("6. XSS Detection")
    xss = all_results.get("xss", {})
    if not xss.get("has_params"):
        add_para("No query parameters in URL. Skipping.", r=100, g=100, b=100)
    elif not xss.get("vulnerable_params"):
        add_para("[OK] No reflected XSS detected.", bold=True, r=34, g=197, b=94)
    else:
        add_table(
            ["Parameter", "Payload"],
            [[v["param"], v["payload"]] for v in xss["vulnerable_params"]],
            col_widths=[4, 12.5]
        )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 7. open redirect
    add_heading("7. Open Redirect Detection")
    redir = all_results.get("open_redirect", {})
    if not redir.get("vulnerable_params"):
        add_para("[OK] No open redirect vulnerabilities detected.", bold=True, r=34, g=197, b=94)
    else:
        add_table(
            ["Parameter", "Payload"],
            [[v["param"], v["payload"]] for v in redir["vulnerable_params"]],
            col_widths=[4, 12.5]
        )
    doc.add_paragraph()

    # ------------------------------------------------------------------ 8. summary
    add_heading("8. Scan Summary")
    if not findings:
        add_para("[OK] No significant vulnerabilities detected!", bold=True, r=34, g=197, b=94)
    else:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        for i, h in enumerate(["Severity", "Finding"]):
            cell = t.rows[0].cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            shade_cell(cell, "1e3a5f")
        for sev, desc in findings:
            row = t.add_row()
            sc = row.cells[0]
            sc.text = sev
            rr, rg, rb = SEV_COLORS.get(sev, (0, 0, 0))
            run = sc.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(rr, rg, rb)
            dc = row.cells[1]
            dc.text = desc
            dc.paragraphs[0].runs[0].font.size = Pt(9)
        for row in t.rows:
            row.cells[0].width = Cm(3)
            row.cells[1].width = Cm(13.5)
        doc.add_paragraph()
        add_para(f"Total findings: {len(findings)}", bold=True, r=220, g=38, b=38, size=11)

    # ------------------------------------------------------------------ footer
    doc.add_paragraph()
    doc.add_paragraph("-" * 80)
    footer = doc.add_paragraph(
        "Generated by WebVulnScanner v1.0 - For authorized security testing only."
    )
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    doc.save(path)
    return path
