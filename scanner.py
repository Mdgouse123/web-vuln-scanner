"""
scanner.py - Main CLI entry point for the Web Vulnerability Scanner.
Usage: python scanner.py <url> [--report json|html|both]
"""

import argparse
import json
import sys
import os

# Fix Windows terminal encoding so rich output renders correctly
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
from datetime import datetime
from urllib.parse import urlparse

from rich.console import Console
from rich.panel import Panel
from rich.table import Table
from rich import box
from rich.rule import Rule
from rich.progress import Progress, SpinnerColumn, TextColumn

from modules.core import get_base_info, check_ssl, normalize_url
from modules.headers import check_security_headers
from modules.sensitive_files import check_sensitive_files
from modules.sqli import check_sqli
from modules.xss import check_xss
from modules.open_redirect import check_open_redirect
from modules.report_docx import export_docx as _export_docx

console = Console()

SEVERITY_COLORS = {
    "CRITICAL": "bold red",
    "HIGH": "red",
    "MEDIUM": "yellow",
    "LOW": "cyan",
    "INFO": "dim white",
}

SEVERITY_ICONS = {
    "CRITICAL": "[bold red][X] CRITICAL[/bold red]",
    "HIGH": "[red][!] HIGH[/red]",
    "MEDIUM": "[yellow][~] MEDIUM[/yellow]",
    "LOW": "[cyan][>] LOW[/cyan]",
    "INFO": "[dim][i] INFO[/dim]",
}

SEVERITY_ORDER = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3, "INFO": 4}


# ---------------------------------------------------------------------------
# Banner
# ---------------------------------------------------------------------------

def print_banner():
    console.print(Panel(
        "[bold cyan]Web Vulnerability Scanner[/bold cyan]  [dim]v1.0[/dim]\n"
        "[dim]For authorized security testing only[/dim]",
        border_style="cyan",
        expand=False,
    ))
    console.print()


# ---------------------------------------------------------------------------
# Section helpers
# ---------------------------------------------------------------------------

def section(title):
    console.print()
    console.print(Rule(f"[bold white]{title}[/bold white]", style="cyan"))


def badge(severity):
    return SEVERITY_ICONS.get(severity, severity)


# ---------------------------------------------------------------------------
# Per-module printers
# ---------------------------------------------------------------------------

def print_base_info(info):
    section("Target Info")
    if "error" in info:
        console.print(f"  [red]Error:[/red] {info['error']}")
        return
    t = Table(box=box.SIMPLE, show_header=False, padding=(0, 2))
    t.add_column("Key", style="dim", width=20)
    t.add_column("Value", style="white")
    t.add_row("URL", info["url"])
    t.add_row("Final URL", info["final_url"])
    t.add_row("Status Code", str(info["status_code"]))
    t.add_row("Server", info["server"])
    t.add_row("Content-Type", info["content_type"])
    t.add_row("Response Time", f"{info['response_time_ms']} ms")
    if len(info["redirect_chain"]) > 1:
        t.add_row("Redirects", " â†’ ".join(info["redirect_chain"]))
    console.print(t)


def print_ssl_info(ssl_info):
    section("SSL / TLS Certificate")
    if ssl_info.get("error"):
        console.print(f"  [red]âœ– SSL Error:[/red] {ssl_info['error']}")
        return
    days = ssl_info.get("days_remaining")
    if days is None:
        console.print("  [dim]Not applicable (HTTP).[/dim]")
        return
    color = "green" if days > 30 else ("yellow" if days > 7 else "red")
    console.print(f"  [green]âœ”[/green] Certificate valid")
    console.print(f"  Issuer:  [white]{ssl_info['ssl_issuer']}[/white]")
    console.print(
        f"  Expires: [{color}]{ssl_info['ssl_expiry']}[/{color}]"
        f"  ([{color}]{days} days remaining[/{color}])"
    )


def print_headers_report(result):
    section("Security Headers")
    if "error" in result:
        console.print(f"  [red]Error:[/red] {result['error']}")
        return
    score, max_score = result["score"], result["max_score"]
    pct = int((score / max_score) * 100) if max_score else 0
    sc = "green" if pct >= 70 else ("yellow" if pct >= 40 else "red")
    console.print(f"  Header Score: [{sc}]{score}/{max_score}[/{sc}]  ({pct}%)\n")

    if result["missing"]:
        console.print("  [bold red]Missing Headers:[/bold red]")
        for h in result["missing"]:
            console.print(f"    {badge(h['severity'])}  [white]{h['header']}[/white]")
            console.print(f"      [dim]{h['description']}[/dim]")
            console.print(f"      [green]Fix:[/green] {h['recommendation']}\n")

    if result["present"]:
        console.print("  [bold green]Present Headers:[/bold green]")
        for h in result["present"]:
            console.print(
                f"    [green]âœ”[/green] [white]{h['header']}[/white]: "
                f"[dim]{h['value'][:80]}[/dim]"
            )

    if result["discouraged"]:
        console.print()
        console.print("  [bold yellow]Information Disclosure Headers:[/bold yellow]")
        for h in result["discouraged"]:
            console.print(f"    {badge(h['severity'])}  [white]{h['header']}[/white]: [dim]{h['value']}[/dim]")
            console.print(f"      [dim]{h['description']}[/dim]")
            console.print(f"      [green]Fix:[/green] {h['recommendation']}")


def print_sensitive_files_report(result):
    section("Sensitive Files & Directories")
    if result.get("error"):
        console.print(f"  [red]Error:[/red] {result['error']}")
        return
    found = result["found"]
    color = "red" if found else "green"
    console.print(
        f"  Probed [white]{result['checked']}[/white] paths â€” "
        f"found [{color}]{len(found)}[/{color}] exposed.\n"
    )
    if not found:
        console.print("  [green]âœ” No sensitive files exposed.[/green]")
        return
    for item in found:
        sc_color = "green" if item["status_code"] == 200 else "yellow"
        console.print(
            f"  {badge(item['severity'])}  [white]{item['path']}[/white]  "
            f"[[{sc_color}]{item['status_code']}[/{sc_color}]]"
        )
        console.print(f"    [dim]{item['description']}[/dim]")
        if item.get("snippet"):
            console.print(f"    [dim]Preview: {item['snippet'][:100]}[/dim]")
        console.print()


def print_sqli_report(result):
    section("SQL Injection Detection")
    if result.get("error"):
        console.print(f"  [red]Error:[/red] {result['error']}")
        return
    if not result["has_params"]:
        console.print("  [dim]No query parameters in URL. Skipping.[/dim]")
        console.print("  [dim]Tip: Try https://example.com/page?id=1[/dim]")
        return
    console.print(f"  Tested params: [white]{', '.join(result['tested_params'])}[/white]\n")
    if not result["vulnerable_params"]:
        console.print("  [green]âœ” No SQL injection errors detected.[/green]")
        return
    console.print("  [bold red]Potential SQL Injection Found:[/bold red]")
    for v in result["vulnerable_params"]:
        console.print(f"    [red]âœ–[/red] Parameter: [white]{v['param']}[/white]")
        console.print(f"      Payload:   [dim]{v['payload']}[/dim]")
        console.print(f"      Triggered: [dim]{v['matched_signature']}[/dim]")
        console.print(f"      URL:       [dim]{v['injected_url'][:100]}[/dim]\n")


def print_xss_report(result):
    section("XSS (Cross-Site Scripting) Detection")
    if result.get("error"):
        console.print(f"  [red]Error:[/red] {result['error']}")
        return
    if not result["has_params"]:
        console.print("  [dim]No query parameters in URL. Skipping.[/dim]")
        console.print("  [dim]Tip: Try https://example.com/search?q=test[/dim]")
        return
    console.print(f"  Tested params: [white]{', '.join(result['tested_params'])}[/white]\n")
    if not result["vulnerable_params"]:
        console.print("  [green]âœ” No reflected XSS detected.[/green]")
        return
    console.print("  [bold red]Potential Reflected XSS Found:[/bold red]")
    for v in result["vulnerable_params"]:
        console.print(f"    [red]âœ–[/red] Parameter: [white]{v['param']}[/white]")
        console.print(f"      Payload: [dim]{v['payload']}[/dim]")
        console.print(f"      URL:     [dim]{v['injected_url'][:100]}[/dim]\n")


def print_redirect_report(result):
    section("Open Redirect Detection")
    if result.get("error"):
        console.print(f"  [red]Error:[/red] {result['error']}")
        return
    if not result["vulnerable_params"]:
        console.print("  [green]âœ” No open redirect vulnerabilities detected.[/green]")
        return
    console.print("  [bold red]Open Redirect Found:[/bold red]")
    for v in result["vulnerable_params"]:
        console.print(f"    [red]âœ–[/red] Parameter: [white]{v['param']}[/white]")
        console.print(f"      Payload: [dim]{v['payload']}[/dim]")
        console.print(f"      URL:     [dim]{v['injected_url'][:100]}[/dim]\n")


# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------

def build_findings(all_results):
    findings = []

    for h in all_results.get("headers", {}).get("missing", []):
        findings.append((h["severity"], f"Missing header: {h['header']}"))
    for h in all_results.get("headers", {}).get("discouraged", []):
        findings.append((h["severity"], f"Info-disclosing header: {h['header']}: {h['value']}"))

    for f in all_results.get("sensitive_files", {}).get("found", []):
        findings.append((f["severity"], f"Exposed file: {f['path']} [{f['status_code']}]"))

    for v in all_results.get("sqli", {}).get("vulnerable_params", []):
        findings.append(("HIGH", f"SQL Injection â€” param: {v['param']}"))

    for v in all_results.get("xss", {}).get("vulnerable_params", []):
        findings.append(("HIGH", f"Reflected XSS â€” param: {v['param']}"))

    for v in all_results.get("open_redirect", {}).get("vulnerable_params", []):
        findings.append(("MEDIUM", f"Open Redirect â€” param: {v['param']}"))

    ssl_info = all_results.get("ssl", {})
    if ssl_info.get("error"):
        findings.append(("HIGH", f"SSL issue: {ssl_info['error']}"))
    elif ssl_info.get("days_remaining") is not None and ssl_info["days_remaining"] < 30:
        findings.append(("MEDIUM", f"SSL cert expiring in {ssl_info['days_remaining']} days"))

    findings.sort(key=lambda x: SEVERITY_ORDER.get(x[0], 5))
    return findings


def print_summary(all_results):
    section("Scan Summary")
    findings = build_findings(all_results)

    if not findings:
        console.print("  [bold green]âœ” No significant vulnerabilities detected![/bold green]")
    else:
        t = Table(box=box.SIMPLE_HEAD, show_header=True, header_style="bold white")
        t.add_column("Severity", width=12)
        t.add_column("Finding")
        for sev, desc in findings:
            color = SEVERITY_COLORS.get(sev, "white")
            t.add_row(f"[{color}]{sev}[/{color}]", desc)
        console.print(t)
        console.print(f"\n  Total findings: [bold red]{len(findings)}[/bold red]")

    console.print(f"\n  [dim]Scan completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}[/dim]")
    return findings


# ---------------------------------------------------------------------------
# Report export
# ---------------------------------------------------------------------------

def export_json(url, all_results, findings):
    os.makedirs("reports", exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    host = urlparse(normalize_url(url)).hostname or "unknown"
    path = f"reports/scan_{host}_{ts}.json"
    payload = {
        "scan_target": url,
        "scan_time": datetime.now().isoformat(),
        "summary": [{"severity": s, "finding": d} for s, d in findings],
        "details": all_results,
    }
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, default=str)
    console.print(f"\n  [green]âœ” JSON report:[/green] [white]{path}[/white]")


def export_pdf(url, all_results, findings):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )

    os.makedirs("reports", exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    host = urlparse(normalize_url(url)).hostname or "unknown"
    path = f"reports/scan_{host}_{ts}.pdf"

    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    style_title   = ParagraphStyle("title",   fontSize=20, textColor=colors.HexColor("#38bdf8"), spaceAfter=6, fontName="Helvetica-Bold")
    style_sub     = ParagraphStyle("sub",     fontSize=9,  textColor=colors.HexColor("#94a3b8"), spaceAfter=14)
    style_h2      = ParagraphStyle("h2",      fontSize=13, textColor=colors.HexColor("#7dd3fc"), spaceBefore=16, spaceAfter=6, fontName="Helvetica-Bold")
    style_body    = ParagraphStyle("body",    fontSize=9,  textColor=colors.HexColor("#e2e8f0"), spaceAfter=4)
    style_label   = ParagraphStyle("label",   fontSize=9,  textColor=colors.HexColor("#94a3b8"))
    style_ok      = ParagraphStyle("ok",      fontSize=9,  textColor=colors.HexColor("#22c55e"), spaceAfter=4)
    style_warn    = ParagraphStyle("warn",    fontSize=9,  textColor=colors.HexColor("#facc15"), spaceAfter=4)
    style_danger  = ParagraphStyle("danger",  fontSize=9,  textColor=colors.HexColor("#ef4444"), spaceAfter=4)
    style_footer  = ParagraphStyle("footer",  fontSize=8,  textColor=colors.HexColor("#475569"), spaceBefore=20)

    BG = colors.HexColor("#0f172a")
    SEV_COLORS = {
        "CRITICAL": colors.HexColor("#dc2626"),
        "HIGH":     colors.HexColor("#ea580c"),
        "MEDIUM":   colors.HexColor("#d97706"),
        "LOW":      colors.HexColor("#0891b2"),
        "INFO":     colors.HexColor("#6b7280"),
    }

    story = []
    base = all_results.get("base_info", {})
    ssl_info = all_results.get("ssl", {})

    # --- Title ---
    story.append(Paragraph("Web Vulnerability Scan Report", style_title))
    story.append(Paragraph(
        f"Target: <b>{url}</b> &nbsp;&nbsp; | &nbsp;&nbsp; "
        f"Scan Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} &nbsp;&nbsp; | &nbsp;&nbsp; "
        f"Status: {base.get('status_code','N/A')} &nbsp;&nbsp; "
        f"Server: {base.get('server','N/A')} &nbsp;&nbsp; "
        f"Response: {base.get('response_time_ms','N/A')} ms",
        style_sub
    ))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1e3a5f")))

    # --- Target Info ---
    story.append(Paragraph("Target Info", style_h2))
    info_data = [
        ["URL", base.get("url", "N/A")],
        ["Final URL", base.get("final_url", "N/A")],
        ["Status Code", str(base.get("status_code", "N/A"))],
        ["Server", base.get("server", "N/A")],
        ["Content-Type", base.get("content_type", "N/A")],
        ["Response Time", f"{base.get('response_time_ms','N/A')} ms"],
    ]
    t_info = Table(info_data, colWidths=[4*cm, 13*cm])
    t_info.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,-1), colors.HexColor("#0f172a")),
        ("TEXTCOLOR",   (0,0), (0,-1),  colors.HexColor("#94a3b8")),
        ("TEXTCOLOR",   (1,0), (1,-1),  colors.HexColor("#e2e8f0")),
        ("FONTSIZE",    (0,0), (-1,-1), 9),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [colors.HexColor("#0f172a"), colors.HexColor("#1e293b")]),
        ("GRID",        (0,0), (-1,-1), 0.3, colors.HexColor("#1e3a5f")),
        ("PADDING",     (0,0), (-1,-1), 5),
    ]))
    story.append(t_info)

    # --- SSL ---
    story.append(Paragraph("SSL / TLS Certificate", style_h2))
    if ssl_info.get("error"):
        story.append(Paragraph(f"[X] SSL Error: {ssl_info['error']}", style_danger))
    elif ssl_info.get("days_remaining") is None:
        story.append(Paragraph("Not applicable (HTTP).", style_label))
    else:
        days = ssl_info["days_remaining"]
        day_style = style_ok if days > 30 else (style_warn if days > 7 else style_danger)
        story.append(Paragraph(f"[OK] Certificate valid  |  Issuer: {ssl_info['ssl_issuer']}", style_ok))
        story.append(Paragraph(f"Expires: {ssl_info['ssl_expiry']}  ({days} days remaining)", day_style))

    # --- Security Headers ---
    story.append(Paragraph("Security Headers", style_h2))
    hr = all_results.get("headers", {})
    if "error" not in hr:
        score, max_score = hr.get("score", 0), hr.get("max_score", 1)
        pct = int((score / max_score) * 100)
        sc_style = style_ok if pct >= 70 else (style_warn if pct >= 40 else style_danger)
        story.append(Paragraph(f"Header Score: {score}/{max_score}  ({pct}%)", sc_style))

        missing = hr.get("missing", [])
        if missing:
            story.append(Paragraph("Missing Headers:", style_warn))
            mdata = [["Severity", "Header", "Recommendation"]]
            for h in missing:
                mdata.append([h["severity"], h["header"], h["recommendation"]])
            mt = Table(mdata, colWidths=[2.5*cm, 5*cm, 9.5*cm])
            mt.setStyle(TableStyle([
                ("BACKGROUND",  (0,0), (-1,0),  colors.HexColor("#1e3a5f")),
                ("TEXTCOLOR",   (0,0), (-1,0),  colors.HexColor("#bae6fd")),
                ("FONTSIZE",    (0,0), (-1,-1), 8),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.HexColor("#0f172a"), colors.HexColor("#1e293b")]),
                ("TEXTCOLOR",   (0,1), (-1,-1), colors.HexColor("#e2e8f0")),
                ("GRID",        (0,0), (-1,-1), 0.3, colors.HexColor("#1e3a5f")),
                ("PADDING",     (0,0), (-1,-1), 4),
                ("FONTNAME",    (0,0), (-1,0),  "Helvetica-Bold"),
            ]))
            story.append(mt)

        present = hr.get("present", [])
        if present:
            story.append(Spacer(1, 6))
            story.append(Paragraph("Present Headers:", style_ok))
            for h in present:
                story.append(Paragraph(f"[OK] {h['header']}: {h['value'][:80]}", style_body))

        disc = hr.get("discouraged", [])
        if disc:
            story.append(Spacer(1, 6))
            story.append(Paragraph("Information Disclosure Headers:", style_warn))
            for h in disc:
                story.append(Paragraph(f"[!] {h['header']}: {h['value']}  —  {h['recommendation']}", style_warn))

    # --- Sensitive Files ---
    story.append(Paragraph("Sensitive Files & Directories", style_h2))
    sf = all_results.get("sensitive_files", {})
    found_files = sf.get("found", [])
    story.append(Paragraph(
        f"Probed {sf.get('checked', 0)} paths — found {len(found_files)} exposed.",
        style_danger if found_files else style_ok
    ))
    if found_files:
        fdata = [["Severity", "Path", "Status", "Description"]]
        for item in found_files:
            fdata.append([item["severity"], item["path"], str(item["status_code"]), item["description"][:60]])
        ft = Table(fdata, colWidths=[2.2*cm, 4*cm, 1.8*cm, 9*cm])
        ft.setStyle(TableStyle([
            ("BACKGROUND",  (0,0), (-1,0),  colors.HexColor("#1e3a5f")),
            ("TEXTCOLOR",   (0,0), (-1,0),  colors.HexColor("#bae6fd")),
            ("FONTSIZE",    (0,0), (-1,-1), 8),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.HexColor("#0f172a"), colors.HexColor("#1e293b")]),
            ("TEXTCOLOR",   (0,1), (-1,-1), colors.HexColor("#e2e8f0")),
            ("GRID",        (0,0), (-1,-1), 0.3, colors.HexColor("#1e3a5f")),
            ("PADDING",     (0,0), (-1,-1), 4),
            ("FONTNAME",    (0,0), (-1,0),  "Helvetica-Bold"),
        ]))
        story.append(ft)

    # --- SQLi ---
    story.append(Paragraph("SQL Injection Detection", style_h2))
    sqli = all_results.get("sqli", {})
    if not sqli.get("has_params"):
        story.append(Paragraph("No query parameters in URL. Skipping.", style_label))
    elif not sqli.get("vulnerable_params"):
        story.append(Paragraph("[OK] No SQL injection errors detected.", style_ok))
    else:
        for v in sqli["vulnerable_params"]:
            story.append(Paragraph(f"[X] Parameter: {v['param']}  |  Payload: {v['payload']}  |  Signature: {v['matched_signature']}", style_danger))

    # --- XSS ---
    story.append(Paragraph("XSS Detection", style_h2))
    xss = all_results.get("xss", {})
    if not xss.get("has_params"):
        story.append(Paragraph("No query parameters in URL. Skipping.", style_label))
    elif not xss.get("vulnerable_params"):
        story.append(Paragraph("[OK] No reflected XSS detected.", style_ok))
    else:
        for v in xss["vulnerable_params"]:
            story.append(Paragraph(f"[X] Parameter: {v['param']}  |  Payload: {v['payload']}", style_danger))

    # --- Open Redirect ---
    story.append(Paragraph("Open Redirect Detection", style_h2))
    redir = all_results.get("open_redirect", {})
    if not redir.get("vulnerable_params"):
        story.append(Paragraph("[OK] No open redirect vulnerabilities detected.", style_ok))
    else:
        for v in redir["vulnerable_params"]:
            story.append(Paragraph(f"[X] Parameter: {v['param']}  |  Payload: {v['payload']}", style_danger))

    # --- Summary ---
    story.append(Paragraph("Scan Summary", style_h2))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1e3a5f")))
    if not findings:
        story.append(Paragraph("[OK] No significant vulnerabilities detected!", style_ok))
    else:
        sdata = [["Severity", "Finding"]]
        for sev, desc in findings:
            sdata.append([sev, desc])
        st = Table(sdata, colWidths=[2.5*cm, 14.5*cm])
        st.setStyle(TableStyle([
            ("BACKGROUND",  (0,0), (-1,0),  colors.HexColor("#1e3a5f")),
            ("TEXTCOLOR",   (0,0), (-1,0),  colors.HexColor("#bae6fd")),
            ("FONTSIZE",    (0,0), (-1,-1), 9),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.HexColor("#0f172a"), colors.HexColor("#1e293b")]),
            ("GRID",        (0,0), (-1,-1), 0.3, colors.HexColor("#1e3a5f")),
            ("PADDING",     (0,0), (-1,-1), 5),
            ("FONTNAME",    (0,0), (-1,0),  "Helvetica-Bold"),
        ]))
        # Color severity column per row
        for i, (sev, _) in enumerate(findings, start=1):
            st.setStyle(TableStyle([
                ("TEXTCOLOR", (0,i), (0,i), SEV_COLORS.get(sev, colors.white)),
                ("TEXTCOLOR", (1,i), (1,i), colors.HexColor("#e2e8f0")),
                ("FONTNAME",  (0,i), (0,i), "Helvetica-Bold"),
            ]))
        story.append(st)
        story.append(Paragraph(f"Total findings: {len(findings)}", style_danger))

    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#1e3a5f")))
    story.append(Paragraph(
        "Generated by WebVulnScanner v1.0 — For authorized security testing only.",
        style_footer
    ))

    # Build with dark background on every page
    def dark_bg(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#0f172a"))
        canvas.rect(0, 0, A4[0], A4[1], fill=1, stroke=0)
        canvas.restoreState()

    doc.build(story, onFirstPage=dark_bg, onLaterPages=dark_bg)
    console.print(f"  [green][OK] PDF report:[/green] [white]{path}[/white]")


def export_docx(url, all_results, findings):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    os.makedirs("reports", exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    host = urlparse(normalize_url(url)).hostname or "unknown"
    path = f"reports/scan_{host}_{ts}.docx"

    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # --- Helper: colored heading ---
    def add_heading(text, level=1, r=41, g=128, b=185):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(r, g, b)
        return p

    # --- Helper: colored paragraph ---
    def add_para(text, bold=False, r=0, g=0, b=0, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(r, g, b)
        return p

    # --- Helper: add a table with header row ---
    def add_table(headers, rows, col_widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        # Header row
        hrow = t.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            # Dark blue background
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "1e3a5f")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
        # Data rows
        for row_data in rows:
            row = t.add_row()
            for i, val in enumerate(row_data):
                cell = row.cells[i]
                cell.text = str(val)
                cell.paragraphs[0].runs[0].font.size = Pt(9)
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        return t

    SEV_COLORS = {
        "CRITICAL": (220, 38,  38),
        "HIGH":     (234, 88,  12),
        "MEDIUM":   (217, 119, 6),
        "LOW":      (8,   145, 178),
        "INFO":     (107, 114, 128),
    }

    base     = all_results.get("base_info", {})
    ssl_info = all_results.get("ssl", {})

    # ===== TITLE =====
    title = doc.add_heading("Web Vulnerability Scan Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 58, 95)

    doc.add_paragraph(
        f"Target: {url}  |  Scan Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  "
        f"Status: {base.get('status_code','N/A')}  |  Server: {base.get('server','N/A')}"
    ).runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ===== TARGET INFO =====
    add_heading("1. Target Info", level=1)
    add_table(
        ["Field", "Value"],
        [
            ["URL",           base.get("url", "N/A")],
            ["Final URL",     base.get("final_url", "N/A")],
            ["Status Code",   str(base.get("status_code", "N/A"))],
            ["Server",        base.get("server", "N/A")],
            ["Content-Type",  base.get("content_type", "N/A")],
            ["Response Time", f"{base.get('response_time_ms','N/A')} ms"],
        ],
        col_widths=[4, 12]
    )
    doc.add_paragraph()

    # ===== SSL =====
    add_heading("2. SSL / TLS Certificate", level=1)
    if ssl_info.get("error"):
        add_para(f"[X] SSL Error: {ssl_info['error']}", bold=True, r=220, g=38, b=38)
    elif ssl_info.get("days_remaining") is None:
        add_para("Not applicable (HTTP).", r=100, g=100, b=100)
    else:
        days = ssl_info["days_remaining"]
        cr, cg, cb = (34, 197, 94) if days > 30 else ((250, 204, 21) if days > 7 else (220, 38, 38))
        add_para(f"[OK] Certificate valid  |  Issuer: {ssl_info['ssl_issuer']}", bold=True, r=34, g=197, b=94)
        add_para(f"Expires: {ssl_info['ssl_expiry']}  ({days} days remaining)", r=cr, g=cg, b=cb)
    doc.add_paragraph()

    # ===== SECURITY HEADERS =====
    add_heading("3. Security Headers", level=1)
    hr = all_results.get("headers", {})
    if "error" not in hr:
        score, max_score = hr.get("score", 0), hr.get("max_score", 1)
        pct = int((score / max_score) * 100)
        cr, cg, cb = (34, 197, 94) if pct >= 70 else ((250, 204, 21) if pct >= 40 else (220, 38, 38))
        add_para(f"Header Score: {score}/{max_score}  ({pct}%)", bold=True, r=cr, g=cg, b=cb)

        missing = hr.get("missing", [])
        if missing:
            add_para("Missing Headers:", bold=True, r=220, g=38, b=38)
            add_table(
                ["Severity", "Header", "Recommendation"],
                [[h["severity"], h["header"], h["recommendation"]] for h in missing],
                col_widths=[2.5, 5, 9]
            )
            doc.add_paragraph()

        present = hr.get("present", [])
        if present:
            add_para("Present Headers:", bold=True, r=34, g=197, b=94)
            add_table(
                ["Header", "Value"],
                [[h["header"], h["value"][:80]] for h in present],
                col_widths=[5, 11.5]
            )
            doc.add_paragraph()

        disc = hr.get("discouraged", [])
        if disc:
            add_para("Information Disclosure Headers:", bold=True, r=250, g=204, b=21)
            add_table(
                ["Header", "Value", "Recommendation"],
                [[h["header"], h["value"], h["recommendation"]] for h in disc],
                col_widths=[3, 4, 9.5]
            )
    doc.add_paragraph()

    # ===== SENSITIVE FILES =====
    add_heading("4. Sensitive Files & Directories", level=1)
    sf = all_results.get("sensitive_files", {})
    found_files = sf.get("found", [])
    cr, cg, cb = (220, 38, 38) if found_files else (34, 197, 94)
    add_para(
        f"Probed {sf.get('checked', 0)} paths - found {len(found_files)} exposed.",
        bold=True, r=cr, g=cg, b=cb
    )
    if found_files:
        add_table(
            ["Severity", "Path", "Status", "Description"],
            [[f["severity"], f["path"], str(f["status_code"]), f["description"][:60]] for f in found_files],
            col_widths=[2.2, 4, 1.8, 8.5]
        )
    doc.add_paragraph()

    # ===== SQLi =====
    add_heading("5. SQL Injection Detection", level=1)
    sqli = all_results.get("sqli", {})
    if not sqli.get("has_params"):
        add_para("No query parameters in URL. Skipping.", r=100, g=100, b=100)
    elif not sqli.get("vulnerable_params"):
        add_para("[OK] No SQL injection errors detected.", bold=True, r=34, g=197, b=94)
    else:
        add_table(
            ["Parameter", "Payload", "Matched Signature"],
            [[v["param"], v["payload"], v["matched_signature"]] for v in sqli["vulnerable_params"]],
            col_widths=[3, 6, 7.5]
        )
    doc.add_paragraph()

    # ===== XSS =====
    add_heading("6. XSS Detection", level=1)
    xss = all_results.get("xss", {})
    if not xss.get("has_params"):
        add_para("No query parameters in URL. Skipping.", r=100, g=100, b=100)
    elif not xss.get("vulnerable_params"):
        add_para("[OK] No reflected XSS detected.", bold=True, r=34, g=197, b=94)
    else:
        add_table(
            ["Parameter", "Payload"],
            [[v["param"], v["payload"]] for v in xss["vulnerable_params"]],
            col_widths=[4, 12.5]
        )
    doc.add_paragraph()

    # ===== OPEN REDIRECT =====
    add_heading("7. Open Redirect Detection", level=1)
    redir = all_results.get("open_redirect", {})
    if not redir.get("vulnerable_params"):
        add_para("[OK] No open redirect vulnerabilities detected.", bold=True, r=34, g=197, b=94)
    else:
        add_table(
            ["Parameter", "Payload"],
            [[v["param"], v["payload"]] for v in redir["vulnerable_params"]],
            col_widths=[4, 12.5]
        )
    doc.add_paragraph()

    # ===== SUMMARY =====
    add_heading("8. Scan Summary", level=1)
    if not findings:
        add_para("[OK] No significant vulnerabilities detected!", bold=True, r=34, g=197, b=94)
    else:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        hrow = t.rows[0]
        for i, h in enumerate(["Severity", "Finding"]):
            cell = hrow.cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "1e3a5f")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
        for sev, desc in findings:
            row = t.add_row()
            # Severity cell with color
            sev_cell = row.cells[0]
            sev_cell.text = sev
            r2, g2, b2 = SEV_COLORS.get(sev, (0, 0, 0))
            sev_run = sev_cell.paragraphs[0].runs[0]
            sev_run.bold = True
            sev_run.font.size = Pt(9)
            sev_run.font.color.rgb = RGBColor(r2, g2, b2)
            # Finding cell
            desc_cell = row.cells[1]
            desc_cell.text = desc
            desc_cell.paragraphs[0].runs[0].font.size = Pt(9)
        # Column widths
        for row in t.rows:
            row.cells[0].width = Cm(3)
            row.cells[1].width = Cm(13.5)

        doc.add_paragraph()
        add_para(f"Total findings: {len(findings)}", bold=True, r=220, g=38, b=38, size=11)

    # ===== FOOTER =====
    doc.add_paragraph()
    doc.add_paragraph("_" * 80)
    footer_p = doc.add_paragraph("Generated by WebVulnScanner v1.0 - For authorized security testing only.")
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    doc.save(path)
    console.print(f"  [green][OK] DOCX report:[/green] [white]{path}[/white]")


def export_html(url, all_results, findings):
    os.makedirs("reports", exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    host = urlparse(normalize_url(url)).hostname or "unknown"
    path = f"reports/scan_{host}_{ts}.html"

    sev_colors = {
        "CRITICAL": "#dc2626", "HIGH": "#ea580c",
        "MEDIUM": "#d97706", "LOW": "#0891b2", "INFO": "#6b7280",
    }
    rows = "".join(
        f'<tr><td style="color:{sev_colors.get(s,"#fff")};font-weight:bold">{s}</td>'
        f"<td>{d}</td></tr>\n"
        for s, d in findings
    )
    base = all_results.get("base_info", {})
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>Scan Report â€” {url}</title>
  <style>
    body{{font-family:'Segoe UI',sans-serif;background:#0f172a;color:#e2e8f0;margin:0;padding:24px}}
    h1{{color:#38bdf8}}h2{{color:#7dd3fc;border-bottom:1px solid #1e3a5f;padding-bottom:6px}}
    table{{width:100%;border-collapse:collapse;margin-bottom:24px}}
    th{{background:#1e3a5f;color:#bae6fd;padding:10px;text-align:left}}
    td{{padding:8px 10px;border-bottom:1px solid #1e293b}}
    tr:hover td{{background:#1e293b}}
    .meta{{color:#94a3b8;font-size:.9em;margin-bottom:24px}}
    .footer{{margin-top:40px;color:#475569;font-size:.8em}}
  </style>
</head>
<body>
  <h1>Web Vulnerability Scan Report</h1>
  <div class="meta">
    <strong>Target:</strong> {url}<br>
    <strong>Time:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br>
    <strong>Status:</strong> {base.get('status_code','N/A')} &nbsp;|&nbsp;
    <strong>Server:</strong> {base.get('server','N/A')} &nbsp;|&nbsp;
    <strong>Response:</strong> {base.get('response_time_ms','N/A')} ms
  </div>
  <h2>Summary â€” {len(findings)} Finding(s)</h2>
  <table>
    <tr><th>Severity</th><th>Finding</th></tr>
    {rows or '<tr><td colspan="2" style="color:#22c55e">No significant vulnerabilities detected.</td></tr>'}
  </table>
  <div class="footer">Generated by WebVulnScanner v1.0 â€” For authorized security testing only.</div>
</body>
</html>"""
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    console.print(f"  [green]âœ” HTML report:[/green] [white]{path}[/white]")


# ---------------------------------------------------------------------------
# Main scan orchestrator
# ---------------------------------------------------------------------------

def run_scan(url, report_format=None):
    print_banner()
    url = normalize_url(url)
    console.print(f"  [bold white]Target:[/bold white] [cyan]{url}[/cyan]\n")

    all_results = {}

    with Progress(SpinnerColumn(), TextColumn("[progress.description]{task.description}"),
                  console=console, transient=True) as progress:
        t = progress.add_task("[cyan]Collecting base info...", total=None)
        all_results["base_info"] = get_base_info(url)
        progress.update(t, description="[cyan]Checking SSL certificate...")
        all_results["ssl"] = check_ssl(url)
        progress.update(t, description="[cyan]Analyzing security headers...")
        all_results["headers"] = check_security_headers(url)
        progress.update(t, description="[cyan]Probing sensitive files...")
        all_results["sensitive_files"] = check_sensitive_files(url)
        progress.update(t, description="[cyan]Testing SQL injection...")
        all_results["sqli"] = check_sqli(url)
        progress.update(t, description="[cyan]Testing XSS reflection...")
        all_results["xss"] = check_xss(url)
        progress.update(t, description="[cyan]Testing open redirects...")
        all_results["open_redirect"] = check_open_redirect(url)
        progress.update(t, description="[green]Done.")

    print_base_info(all_results["base_info"])
    print_ssl_info(all_results["ssl"])
    print_headers_report(all_results["headers"])
    print_sensitive_files_report(all_results["sensitive_files"])
    print_sqli_report(all_results["sqli"])
    print_xss_report(all_results["xss"])
    print_redirect_report(all_results["open_redirect"])
    findings = print_summary(all_results)

    if report_format in ("json", "both", "all"):
        export_json(url, all_results, findings)
    if report_format in ("html", "both", "all"):
        export_html(url, all_results, findings)
    if report_format in ("pdf", "all"):
        export_pdf(url, all_results, findings)

    console.print()


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Web Vulnerability Scanner",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python scanner.py https://example.com
  python scanner.py https://example.com --report html
  python scanner.py "https://example.com/page?id=1" --report both

Checks performed:
  - HTTP security headers analysis
  - SSL/TLS certificate validity
  - Sensitive file & directory exposure
  - SQL injection (error-based)
  - Reflected XSS detection
  - Open redirect detection

WARNING: Only use on systems you own or have explicit written permission to test.
        """,
    )
    parser.add_argument("url", help="Target URL (e.g. https://example.com)")
    parser.add_argument(
        "--report",
        choices=["json", "html", "pdf", "both", "all"],
        default=None,
        help="Export results: json, html, pdf, both (html+json), all (html+json+pdf)",
    )
    args = parser.parse_args()

    console.print()
    console.print(
        "[bold yellow]! Legal Notice:[/bold yellow] Only scan systems you own "
        "or have explicit written permission to test.\n"
    )

    try:
        run_scan(args.url, report_format=args.report)
    except KeyboardInterrupt:
        console.print("\n[yellow]Scan interrupted.[/yellow]")
        sys.exit(0)


if __name__ == "__main__":
    main()

